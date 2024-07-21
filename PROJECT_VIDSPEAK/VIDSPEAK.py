from tkinter import *
from PIL import Image, ImageTk
import tkinter.messagebox
import speech_recognition as sr
from google_trans_new import google_translator
from gtts import gTTS
from playsound import playsound
from youtube_transcript_api import YouTubeTranscriptApi
import os
from docx import Document

window = Tk()

def exit1():
    window.destroy()

def about():
    tkinter.messagebox.showerror("ABOUT WINDOW", '''HI SHRIYANSH NEGI (THE DEVELOPER) THIS SIDE,\n
THIS APPLICATION WAS MADE BY ME TO ASSIST THE USER WITH THEIR BASIC TASKS.

SOMETIMES IT BECOMES HECTIC FOR THE USER TO READ ALL THE TEXT MESSAGES, SO MY PROJECT HELPS THOSE USERS TO LISTEN TO THOSE MESSAGES.
\nSOME USERS ARE UNABLE TO MAKE NOTES, SO OUR APPLICATION CAN ALSO DOWNLOAD THE TRANSCRIPT FROM THE VIDEOS AND STORE IT IN A FILE WITH THE NAME ENTERED BY THE USER.''')    


window.geometry("500x600")

AUDIO_LIST = []

def update():
    filename = fl_t.get().strip()
    if filename:
        AUDIO_LIST.append(filename)

def display():
    if not AUDIO_LIST:  # Check if the list is empty
        tkinter.messagebox.showinfo("HISTORY WINDOW", "No files saved yet.")
        return

    files_display = "\n".join(AUDIO_LIST)
    tkinter.messagebox.showinfo("HISTORY WINDOW", f"THE FILES THAT YOU HAVE SAVED TILL NOW ARE:\n{files_display}")

def text_to_voice():
    input_text = fn.get() 
    text = input_text.strip()
    filename = fl_t.get().strip() + ".mp3"
    if not text:
        tkinter.messagebox.showerror("Error", "Input text cannot be empty.")
        return

    try:
        # Create a gTTS object with the input text
        tts = gTTS(text=text, lang="en", slow=False)
        tts.save(filename)  # Save the audio file
        
        # Check if the file exists before trying to play it
        if os.path.exists(filename):
            playsound(filename)  # Play the saved audio file
        else:
            tkinter.messagebox.showerror("Error", "Audio file was not created successfully.")
    except Exception as e:
        tkinter.messagebox.showerror("Error", f"Failed to play audio: {str(e)}")
        
def youtube():
    video_link = jn.get()
    filename = fl.get().strip() + ".docx"  
    
    if not video_link.strip():
        tkinter.messagebox.showerror("ALERT WINDOW", "VIDEO LINK FIELD CAN'T BE EMPTY....\n\nKINDLY FILL THE VIDEO LINK FIELD")
    elif not filename.strip():
        tkinter.messagebox.showerror("ALERT WINDOW", "FILE NAME FIELD CAN'T BE EMPTY...\n\nKINDLY FILL THE FILE NAME")
    else:
        try:
            srt = YouTubeTranscriptApi.get_transcript(video_link)
            
            doc = Document()
            doc.add_heading('Notes', level=1)
            
            for item in srt:
                text = item["text"]
                doc.add_paragraph(text)
                
            doc.save(filename)
            tkinter.messagebox.showinfo("ALERT WINDOW", "SUCCESSFULLY DOWNLOADED !!!!")
        except Exception as e:
            if "Video unavailable" in str(e):
                tkinter.messagebox.showerror("Error", "Video not found. Please enter the correct link.")
            else:
                tkinter.messagebox.showerror("Error", str(e))

def handler():
    if not (fn.get().strip() and fl_t.get().strip()):
        tkinter.messagebox.showerror("ERROR WINDOW", "REQUIRED FIELD IS EMPTY.\n PLEASE TRY AGAIN AFTER ENTRTING FIELDS.")
        
    else:
         text_to_voice()

#CREATING MENU BAR FOR THE APPLICATION
menu = Menu(window)
window.config(menu=menu)

subm1 = Menu(menu)
menu.add_cascade(label="File", menu=subm1)
subm1.add_command(label="Exit", command=exit1)

subm3 = Menu(menu)
menu.add_cascade(label="ABOUT", menu=subm3)
subm3.add_command(label="GAIN KNOWLEDGE FROM HERE!!!", command=about)

subm4 = Menu(menu)
menu.add_cascade(label="HISTORY", menu=subm4)
subm4.add_command(label="RECENT AUDIO FILE SAVED BY YOU :)",command=display)


        
LABEL_11 = Label(window, text="ENTER THE TEXT:", font=("Orbitron", 13, "bold"))
LABEL_11.place(x=10, y=320)

fn = StringVar()
entry_1 = Entry(window, textvariable=fn,width=30,font=("Stencil",10),relief="solid",borderwidth=2)
entry_1.place(x=12, y=350)

LABEL_11 = Label(window, text="FILE NAME HERE:", font=("Orbitron", 13, "bold"))
LABEL_11.place(x=10, y=400)

fl_t = StringVar()
entry_1 = Entry(window, textvariable=fl_t,width=30,font=("Stencil",10),relief="solid",borderwidth=2)
entry_1.place(x=12, y=430)
window.title("VidSpeak APPLICATION")

LABEL_1=Label(text="VidSpeak",font=("Bernard MT Condensed",40,"bold"),fg="white",bg="blue",width=19,relief="solid")
LABEL_1.place(x=8,y=10)

LABEL_2=Label(text="SUNO!\n ARAM SE.....",font=("Lemon",13,"bold"))
LABEL_2.place(x=75,y=120)

LABEL_3=Label(text="YOUSCRIPT \nDOWNLOADER",font=("Lemon",13,"bold"))
LABEL_3.place(x=300,y=120)

# FOR SETTING IMAGE IN OUR GUI INTERFACE

img1 = Image.open("A:\\TKINTER PYTHON for GUI\\MINI PROJECT\\T2.png")
photo1 = ImageTk.PhotoImage(img1)
label_V = Label(window, image=photo1, relief="solid")
label_V.place(x=10, y=180)


button_l=Button(text="LISTEN",font=("Wide Latin",10),command=handler,fg="white",bg="blue",relief="solid",borderwidth=2,height=2)
button_l.place(x=55,y=480)

img2 = Image.open("A:\\TKINTER PYTHON for GUI\\MINI PROJECT\\T1.png")
photo = ImageTk.PhotoImage(img2)
label_Y = Label(window, image=photo, relief="solid")
label_Y.place(x=240, y=180)

label_you=Label(text="PASTE LINK HERE !",font=("Orbitron",13,"bold"))
label_you.place(x=280,y=320)

jn = StringVar()
entry_1 = Entry(window, textvariable=jn, width=27,font=("Stencil",10),relief="solid",borderwidth=2)
entry_1.place(x=283, y=350)

LABEL_4=Label(text="FILE NAME HERE:",font=("Orbitron",13,"bold"))
LABEL_4.place(x=280,y=400)

fl=StringVar()
entry_2=Entry(window,textvar=fl,width=27,font=("Stencil",10),relief="solid",borderwidth=2)
entry_2.place(x=283,y=430)

Button_1=Button(text="DOWNLOAD \nSCRIPT",font=("Wide Latin",10),command=youtube,fg="white",bg="red",relief="solid",borderwidth=2)
Button_1.place(x=288,y=480)
